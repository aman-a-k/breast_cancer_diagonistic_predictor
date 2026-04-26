from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import os

ROOT = Path(__file__).resolve().parents[1]
ARTIFACTS = ROOT / "artifacts"
FIGURES = ARTIFACTS / "figures"

# Use /tmp for all transient files on Vercel
IS_VERCEL = os.environ.get("VERCEL") == "1"
TEMP = Path("/tmp") if IS_VERCEL else (ROOT / "tmp")
TEMP.mkdir(exist_ok=True)

REPORT_DIR = Path("/tmp") if IS_VERCEL else (ROOT / "report")
REPORT_DIR.mkdir(exist_ok=True)

OUTPUT = REPORT_DIR / "Breast_Cancer_Diagnosis_Prediction_Summary.docx"
DETAILED_OUTPUT = REPORT_DIR / "Breast_Cancer_Diagnosis_Prediction_Report.docx"


def add_picture(document: Document, filename: str, caption: str, width_inches: float = 5.8) -> None:
    path = FIGURES / filename
    if not path.exists():
        path = TEMP / filename
    if path.exists():
        document.add_picture(str(path), width=Inches(width_inches))
        p = document.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _safe_float(value: object) -> float | None:
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _generate_patient_plots(sample_snapshot: dict, metrics: dict, prediction: dict, case_id: str) -> list[str]:
    filenames = []
    # Clinical high-contrast styling
    plt.rcParams.update({
        'font.size': 11,
        'axes.labelweight': 'bold',
        'axes.titleweight': 'bold',
        'figure.dpi': 600,
        'axes.spines.right': False,
        'axes.spines.top': False
    })
    
    # 1. Prediction Confidence Plot (EXHIBIT A)
    probs_dict = prediction.get("probabilities", {})
    if probs_dict:
        plt.figure(figsize=(5.5, 3.5))
        sorted_keys = sorted(probs_dict.keys(), key=lambda x: "malignant" not in x.lower())
        labels = [k.title() for k in sorted_keys]
        values = [float(probs_dict[k]) for k in sorted_keys]
        colors = ["#d1495b", "#2a9d8f"]
        
        bars = plt.bar(labels, values, color=colors, alpha=0.9, edgecolor="black", linewidth=1.5)
        plt.ylim(0, 1.2)
        plt.title("Pathological Classification Confidence Cluster", pad=15)
        plt.ylabel("Confidence Level")
        for bar in bars:
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height + 0.02, f'{height*100:.1f}%', ha='center', va='bottom', fontweight='bold', fontsize=12)
        
        fname = f"conf_{case_id}.png"
        plt.savefig(TEMP / fname, dpi=600, bbox_inches='tight')
        plt.close()
        filenames.append(fname)

    # 2. Biological Signal Deviation (EXHIBIT B)
    feature_means = metrics.get("feature_means", {})
    top_importances = metrics.get("feature_importances", [])[:8]
    if feature_means and sample_snapshot:
        plt.figure(figsize=(7, 4.5))
        f_labels, ratios = [], []
        for item in top_importances:
            f = item["feature"]
            avg, val = float(feature_means.get(f, 1)), float(sample_snapshot.get(f, 0))
            ratios.append(val / avg if avg != 0 else 1.0)
            f_labels.append(f.replace(" ", "\n").title())
        
        colors = ["#d1495b" if abs(1.0 - r) > 0.25 else "#2454a6" for r in ratios]
        plt.bar(f_labels, ratios, color=colors, alpha=0.8, edgecolor="black")
        plt.axhline(y=1.0, color="red", linestyle="--", linewidth=2, label="Historical Population Mean")
        plt.title("Biomarker Displacement Profile (Current Patient)", pad=20)
        plt.ylabel("Deviation Ratio (1.0 = Average)")
        plt.grid(axis='y', linestyle=':', alpha=0.6)
        plt.legend(loc='best', frameon=True)
        
        fname = f"ratio_{case_id}.png"
        plt.savefig(TEMP / fname, dpi=600, bbox_inches='tight')
        plt.close()
        filenames.append(fname)

    # 3. Spatial Localization (EXHIBIT C)
    pca_params = metrics.get("cluster_summary", {}).get("projection_params")
    pca_points = metrics.get("cluster_summary", {}).get("pca_points")
    if pca_params and pca_points:
        plt.figure(figsize=(6.5, 5))
        df_pca = pd.DataFrame(pca_points)
        plt.scatter(df_pca[df_pca['label']=='benign']['x'], df_pca[df_pca['label']=='benign']['y'], c="#2a9d8f", alpha=0.1, s=20, label="Historical Benign Area")
        plt.scatter(df_pca[df_pca['label']=='malignant']['x'], df_pca[df_pca['label']=='malignant']['y'], c="#d1495b", alpha=0.1, s=20, label="Historical Malignant Area")
        
        try:
            mean, scale, comp = np.array(pca_params["scaler_mean"]), np.array(pca_params["scaler_scale"]), np.array(pca_params["pca_components"])
            f_names = metrics.get("feature_names", [])
            x_raw = np.array([float(sample_snapshot.get(fn, 0)) for fn in f_names])
            scaled = (x_raw - mean) / scale
            pc1, pc2 = np.dot(scaled, comp[0]), np.dot(scaled, comp[1])
            plt.scatter([pc1], [pc2], c="#f59e0b", s=400, marker="*", edgecolor="black", label="CURRENT PATIENT LOCALIZATION", zorder=100)
        except Exception: pass
            
        plt.title("Case Localization Mapping", pad=20)
        plt.legend(loc='lower left', frameon=True, facecolor='white', framealpha=1.0)
        plt.grid(True, linestyle="--", alpha=0.1)
        
        fname = f"pca_{case_id}.png"
        plt.savefig(TEMP / fname, dpi=600, bbox_inches='tight')
        plt.close()
        filenames.append(fname)
        
    return filenames

def generate_report(
    sample_snapshot: dict | None = None,
    prediction: dict | None = None,
    average_deviation: float | None = None,
    output_path: Path | None = None,
) -> Path:
    metrics = json.loads((ARTIFACTS / "metrics.json").read_text(encoding="utf-8"))
    document = Document()
    case_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    label = prediction.get("prediction_label", "N/A") if prediction else "N/A"
    print(f"DEBUG: Report generating Case {case_id} for Pred: {label}")
    
    header_title = document.add_heading("Diagnostic Pathology Analysis Report", level=0)
    header_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_p = document.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.add_run(f"Case Ref: BC-{case_id}\n").bold = True
    info_p.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    document.add_paragraph("_" * 75).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if prediction:
        document.add_heading("Section I: Diagnostic Conclusion", level=1)
        label = str(prediction.get("prediction_label", "Unknown")).upper()
        p = document.add_paragraph()
        p.add_run("PRIMARY DIAGNOSIS: ").bold = True
        res_run = p.add_run(label)
        res_run.bold = True; res_run.font.size = Pt(18)
        
        if "MALIGNANT" in label:
            res_run.font.color.rgb = RGBColor(0xD1, 0x49, 0x5B)
            document.add_paragraph("WARNING: This profile indicates highly suspicious cellular morphology.", style="Intense Quote")
        else:
            res_run.font.color.rgb = RGBColor(0x2A, 0x9D, 0x8F)
            document.add_paragraph("Interpretation: The diagnostic measurements align with benign historical profiles.", style="Quote")
            
        if sample_snapshot:
            _generate_patient_plots(sample_snapshot, metrics, prediction, case_id)
            add_picture(document, f"conf_{case_id}.png", "Exhibit A: Patient-Specific Diagnostic Probability Breakdown", width_inches=4.5)

    document.add_heading("Section II: Comparative Biological Markers", level=1)
    if sample_snapshot:
        add_picture(document, f"ratio_{case_id}.png", "Exhibit B: Biomarker Shift Profile (Value / Dataset Mean Ratio)", width_inches=5.8)

    document.add_heading("Section III: Cluster Space Mapping", level=1)
    if sample_snapshot:
        add_picture(document, f"pca_{case_id}.png", "Exhibit C: Case localization relative to historical biopsies", width_inches=5.8)

    document.add_heading("Section IV: Quantitative Biomarker Data", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Biomarker'; hdr_cells[1].text = 'Patient Value'; hdr_cells[2].text = 'Predictive Weight'
    
    for item in metrics.get("feature_importances", [])[:7]:
        feat = item["feature"]
        row = table.add_row().cells
        row[0].text = feat.title()
        val = float(sample_snapshot.get(feat, 0)) if sample_snapshot else 0.0
        row[1].text = f"{val:.5f}"; row[2].text = f"{item['importance']*100:.1f}%"

    document.add_heading("Section V: Medical Disclaimer", level=1)
    p = document.add_paragraph("This report is an AI-assisted diagnostic aid. Final medical decisions should be made by a board-certified physician.")
    p.italic = True; p.runs[0].font.size = Pt(8)

    final_filename = f"Diagnostic_Report_{case_id}.docx"
    final_output = REPORT_DIR / final_filename
    document.save(final_output)
    return final_output

def main() -> None:
    output = generate_report(output_path=OUTPUT)
    print(f"Word report written to: {output}")

if __name__ == "__main__":
    main()
